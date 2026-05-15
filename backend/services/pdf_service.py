"""
pdf_service.py — Document Parsing Service
==========================================
Extracts text from PDF and DOCX resume files using:
  - PyMuPDF (fitz) for PDFs
  - python-docx for DOCX files

Also includes basic metadata extraction (pages, word count).
"""

import io
import os
import uuid
import shutil
from pathlib import Path
from typing import Tuple, Optional

from fastapi import UploadFile

from backend.config import get_settings
from backend.utils.logger import get_logger
from backend.utils.validators import sanitize_filename

logger = get_logger(__name__)
settings = get_settings()


# ═══════════════════════════════════════════════════════════════════════════════
# PDF Extraction
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract all text from a PDF using PyMuPDF.

    Args:
        file_bytes: Raw PDF bytes.

    Returns:
        Extracted text as a string.
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")  # Plain text extraction
            pages_text.append(text)

        full_text = "\n".join(pages_text)
        logger.info(f"PDF extracted: {len(doc)} pages, {len(full_text)} chars")
        doc.close()
        return full_text.strip()

    except ImportError:
        logger.error("PyMuPDF (fitz) not installed. Run: pip install pymupdf")
        raise RuntimeError("PDF parsing library not available.")
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}", exc_info=True)
        raise RuntimeError(f"Could not parse PDF: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX Extraction
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from a DOCX file using python-docx.

    Args:
        file_bytes: Raw DOCX bytes.

    Returns:
        Extracted text as a string.
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        full_text = "\n".join(paragraphs)
        logger.info(f"DOCX extracted: {len(paragraphs)} paragraphs, {len(full_text)} chars")
        return full_text.strip()

    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise RuntimeError("DOCX parsing library not available.")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}", exc_info=True)
        raise RuntimeError(f"Could not parse DOCX: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# Main Entry Point
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text_from_file(file_bytes: bytes, file_type: str) -> str:
    """
    Route extraction to the correct parser based on file type.

    Args:
        file_bytes: Raw file bytes.
        file_type: 'pdf' or 'docx'.

    Returns:
        Extracted plain text.
    """
    if file_type == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


# ═══════════════════════════════════════════════════════════════════════════════
# File Storage
# ═══════════════════════════════════════════════════════════════════════════════

async def save_uploaded_file(
    upload_file: UploadFile,
    user_id: int,
    file_type: str
) -> Tuple[str, bytes, int]:
    """
    Save an uploaded file to disk in the uploads directory.

    Args:
        upload_file: FastAPI UploadFile object.
        user_id: ID of the uploading user (for namespacing).
        file_type: 'pdf' or 'docx'.

    Returns:
        Tuple of (stored_path, file_bytes, file_size).
    """
    # Build user-specific upload directory
    user_dir = Path(settings.UPLOAD_DIR) / f"user_{user_id}"
    user_dir.mkdir(parents=True, exist_ok=True)

    # Unique filename to prevent collisions
    safe_name = sanitize_filename(upload_file.filename or "resume")
    unique_name = f"{uuid.uuid4().hex}_{safe_name}"
    dest_path = user_dir / unique_name

    # Read file bytes
    file_bytes = await upload_file.read()
    file_size = len(file_bytes)

    # Check size limit
    if file_size > settings.max_file_size_bytes:
        raise ValueError(
            f"File too large ({file_size // 1024 // 1024}MB). "
            f"Maximum allowed: {settings.MAX_FILE_SIZE_MB}MB"
        )

    # Write to disk
    with open(dest_path, "wb") as f:
        f.write(file_bytes)

    logger.info(f"Saved file: {dest_path} ({file_size} bytes)")
    return str(dest_path), file_bytes, file_size


def delete_file(file_path: str) -> bool:
    """
    Delete a resume file from disk during cleanup.

    Returns:
        True if deleted, False if not found.
    """
    path = Path(file_path)
    if path.exists():
        path.unlink()
        logger.info(f"Deleted file: {file_path}")
        return True
    logger.warning(f"File not found for deletion: {file_path}")
    return False
